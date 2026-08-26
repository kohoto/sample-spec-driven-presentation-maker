# Copyright Amazon.com, Inc. or its affiliates. All Rights Reserved.
# SPDX-License-Identifier: MIT-0
"""Attachment pipeline orchestrator — core read/import logic.

This is the shared core that Local and Remote adapters delegate to.
It receives a materialized local path and performs:
  - read: text projection + paging + guidance header
  - import: conversion + bundle commit

The adapters handle source materialization (path validation / S3 download)
and result formatting (Image content vs path, etc).
"""

from __future__ import annotations

import json
import logging
import shutil
import time
import uuid
from datetime import datetime, timezone
from pathlib import Path
from typing import Any

from sdpm.tools.attachment.bundle import (
    BundleManifest,
    LocalBundleCommitter,
    build_file_list,
    compute_file_sha256,
)
from sdpm.tools.attachment.cache import (
    LocalStageCache,
    StageRecord,
    compute_hash,
    compute_import_key,
    compute_pipeline_key,
    compute_source_identity_hash,
    compute_stage_key,
    get_completed_stages,
    pipeline_version,
    STAGE_ORDER,
)
from sdpm.tools.attachment.errors import (
    ImportConversionFailed,
    ImportConflict,
    ImportIncomplete,
    ImportLimitExceeded,
    SourceLimitExceeded,
    SourceValidationError,
)
from sdpm.tools.attachment.limits import (
    IMPORT_DEADLINE_S,
    MAX_PDF_EXTRACTED_IMAGES_BYTES,
    MAX_PDF_PAGES,
    PAGING_DEFAULT_LIMIT,
)
from sdpm.tools.attachment.paging import PagingError, page_text
from sdpm.tools.attachment.source import (
    detect_media_type,
    enforce_content_limits,
    get_canonical_extension,
    sanitize_filename,
    validate_media_type,
)

logger = logging.getLogger(__name__)


def read_attachment_core(
    materialized_path: Path,
    source: str,
    offset: int = 0,
    limit: int = PAGING_DEFAULT_LIMIT,
    *,
    source_identity: dict[str, Any] | None = None,
    cache: LocalStageCache | None = None,
    extra_header_fields: dict[str, Any] | None = None,
    content_type: str | None = None,
) -> dict[str, Any]:
    """Core read_attachment logic — operates on a materialized local file.

    Args:
        materialized_path: Absolute path to the source file on local disk.
        source: Original source string (for header).
        offset: UTF-8 byte offset.
        limit: Max bytes for response.
        source_identity: Pre-computed identity for cache lookup.
        cache: Stage cache for reusing conversions.
        extra_header_fields: Additional fields for the header (guideInstruction, etc).

    Returns:
        Dict response matching the tool contract.
    """
    if not materialized_path.exists():
        raise SourceValidationError(f"Materialized file not found: {materialized_path}")

    # Detect media type
    with open(materialized_path, "rb") as f:
        header_bytes = f.read(8192)

    file_name = materialized_path.name
    detected_type = detect_media_type(header_bytes, filename=file_name, path=materialized_path)
    media_type = validate_media_type(detected_type, content_type, filename=file_name)
    enforce_content_limits(materialized_path, media_type)

    # Determine kind
    kind = _media_type_to_kind(media_type)

    if kind == "image":
        # Images are handled by adapter (Cloud=Image content, Local=path+color)
        return {
            "kind": "image",
            "source": source,
            "fileName": file_name,
            "mediaType": media_type,
            "path": str(materialized_path),
            "colorAnalysis": _analyze_image_colors(materialized_path, media_type),
        }

    # Get canonical text projection and PPTX workflow metadata.
    header_fields = dict(extra_header_fields or {})
    if kind == "pptx":
        text_data, pptx_fields = _pptx_projection(materialized_path)
        header_fields.update(pptx_fields)
    else:
        text_data = _get_text_projection(materialized_path, media_type, source_identity, cache)

    if text_data is None:
        raise ImportConversionFailed(f"Cannot create text projection for {media_type}")

    # Page the text
    try:
        result = page_text(
            data=text_data,
            offset=offset,
            limit=limit,
            source=source,
            file_name=file_name,
            media_type=media_type,
            kind=kind,
            extra_header_fields=header_fields or None,
        )
    except PagingError as e:
        return {"error": {"code": e.code, "message": e.message}}

    # Build response
    response: dict[str, Any] = {"header": result.header}
    if result.body:
        response["body"] = result.body
    return response


def import_attachment_core(
    materialized_path: Path,
    source: str,
    deck_dir: Path,
    filename: str = "",
    *,
    source_identity: dict[str, Any] | None = None,
    cache: LocalStageCache | None = None,
    content_type: str | None = None,
) -> dict[str, Any]:
    """Convert an attachment and atomically commit a resumable immutable bundle."""
    start_time = time.monotonic()

    if not materialized_path.exists():
        raise SourceValidationError(f"Materialized file not found: {materialized_path}")

    source_hash = compute_file_sha256(materialized_path)
    if not filename:
        filename = materialized_path.name
    filename = sanitize_filename(filename)

    with open(materialized_path, "rb") as file:
        header_bytes = file.read(8192)
    detected_type = detect_media_type(header_bytes, filename=filename, path=materialized_path)
    media_type = validate_media_type(detected_type, content_type, filename=filename)
    enforce_content_limits(materialized_path, media_type)
    canonical_extension = get_canonical_extension(media_type)
    if canonical_extension:
        filename = sanitize_filename(f"{Path(filename).stem}{canonical_extension}")

    version = pipeline_version()
    options = {"filename": filename, "mediaType": media_type}
    import_key = compute_import_key(source_hash, version, options)
    committer = LocalBundleCommitter(deck_dir=deck_dir)
    existing = committer.get_committed(import_key)
    if existing is not None:
        return {
            "importKey": import_key,
            "sourceHash": source_hash,
            "reused": True,
            "files": existing.files,
            "imageMapping": existing.image_mapping,
            "deckJson": existing.deck_json,
            "templatePath": existing.template_path,
            "bundlePath": f"attachments/imports/{import_key}",
        }

    identity = source_identity or {
        "kind": "content",
        "sha256": source_hash,
        "size": materialized_path.stat().st_size,
    }
    source_identity_hash = compute_source_identity_hash(identity)
    options_hash = compute_hash(json.dumps(options, sort_keys=True, separators=(",", ":")))
    pipeline_key = compute_pipeline_key(version, options_hash)
    stage_keys = {
        stage: compute_stage_key(stage, options)
        for stage in STAGE_ORDER
    }

    request_id = str(uuid.uuid4())
    staging = committer.create_staging(request_id)
    completed_stages: list[str] = []

    def _stage_record(stage: str) -> StageRecord:
        return StageRecord(
            stage=stage,
            source_identity_hash=source_identity_hash,
            source_hash=source_hash,
            pipeline_version=version,
            options_hash=options_hash,
            outputs=build_file_list(staging),
            completed_at=_utc_now_iso(),
        )

    def _complete_stage(stage: str) -> None:
        if stage in completed_stages:
            return
        expected_index = len(completed_stages)
        if expected_index >= len(STAGE_ORDER) or STAGE_ORDER[expected_index] != stage:
            raise ImportConversionFailed(f"Attachment stage completed out of order: {stage}")
        if cache is not None:
            cache.publish_stage(
                source_identity_hash,
                pipeline_key,
                stage,
                stage_keys[stage],
                _stage_record(stage),
                staging,
            )
        completed_stages.append(stage)

    def _check_deadline() -> None:
        if time.monotonic() - start_time >= IMPORT_DEADLINE_S:
            raise _DeadlineExceeded(completed_stages)

    try:
        if cache is not None:
            cached_stages = get_completed_stages(
                cache, source_identity_hash, pipeline_key, stage_keys,
            )
            if cached_stages:
                latest_stage = cached_stages[-1]
                latest_record = cache.get_stage(
                    source_identity_hash,
                    pipeline_key,
                    latest_stage,
                    stage_keys[latest_stage],
                )
                if latest_record is not None:
                    cached_outputs = cache.stage_dir(
                        source_identity_hash,
                        pipeline_key,
                        latest_stage,
                        stage_keys[latest_stage],
                    ) / "outputs"
                    shutil.copytree(cached_outputs, staging, dirs_exist_ok=True)
                    completed_stages.extend(cached_stages)

        _check_deadline()
        source_dir = staging / "source"
        source_dir.mkdir(exist_ok=True)
        source_copy = source_dir / filename
        if "materialize" not in completed_stages:
            shutil.copy2(materialized_path, source_copy)
            _complete_stage("materialize")
        elif not source_copy.is_file():
            raise ImportConversionFailed("Cached materialize stage is missing its source output")

        _check_deadline()
        kind = _media_type_to_kind(media_type)
        if "convert_deck" not in completed_stages:
            if kind == "image":
                if "extract_text" not in completed_stages:
                    _complete_stage("extract_text")  # N/A for images
                if "extract_images" not in completed_stages:
                    extracted_images_dir = staging / "extracted" / "images"
                    extracted_images_dir.mkdir(parents=True, exist_ok=True)
                    image_name = f"{source_hash[:8]}_{filename}"
                    destination = extracted_images_dir / image_name
                    if media_type == "image/webp":
                        destination = destination.with_suffix(".png")
                        _convert_webp_to_png(source_copy, destination)
                    else:
                        shutil.copy2(source_copy, destination)
                    _complete_stage("extract_images")
                _complete_stage("convert_deck")
            elif kind == "pptx":
                _import_pptx(source_copy, staging, completed_stages, _complete_stage, _check_deadline)
            else:
                _import_document(
                    source_copy,
                    media_type,
                    staging,
                    completed_stages,
                    _complete_stage,
                    _check_deadline,
                )

        _check_deadline()
        if "validate_bundle" not in completed_stages:
            # Enumerating the bundle also rejects any disappearing output before publication.
            build_file_list(staging)
            _complete_stage("validate_bundle")

        files = build_file_list(staging)
        image_mapping = _build_image_mapping(staging)
        deck_json_path = _find_deck_json(staging)
        template_path = _find_template(staging)
        manifest = BundleManifest(
            import_key=import_key,
            source_hash=source_hash,
            pipeline_version=version,
            options=options,
            files=files,
            image_mapping=image_mapping,
            deck_json=deck_json_path,
            template_path=template_path,
        )
        committer.commit(request_id, manifest)
        return {
            "importKey": import_key,
            "sourceHash": source_hash,
            "reused": False,
            "files": files,
            "imageMapping": image_mapping,
            "deckJson": deck_json_path,
            "templatePath": template_path,
            "bundlePath": f"attachments/imports/{import_key}",
        }
    except _DeadlineExceeded as error:
        committer.cleanup_staging(request_id)
        return ImportIncomplete(
            message="Import did not complete within deadline.",
            completed_stages=error.completed_stages,
        ).to_dict()
    except (ImportLimitExceeded, ImportConversionFailed, ImportConflict, SourceLimitExceeded) as error:
        committer.cleanup_staging(request_id)
        return error.to_dict()
    except Exception as error:
        committer.cleanup_staging(request_id)
        logger.exception("Import failed: %s", error)
        return ImportConversionFailed(str(error)).to_dict()


def _utc_now_iso() -> str:
    """Return an unambiguous UTC timestamp for cache completion records."""
    return datetime.now(timezone.utc).isoformat()


class _DeadlineExceeded(Exception):
    """Internal: deadline reached during import."""

    def __init__(self, completed_stages: list[str]) -> None:
        self.completed_stages = completed_stages



def _analyze_image_colors(path: Path, media_type: str) -> dict[str, Any] | None:
    if media_type == "image/svg+xml":
        return None
    try:
        import colorsys
        from PIL import Image

        with Image.open(path) as image:
            rgb = image.convert("RGB")
            rgb.thumbnail((128, 128))
            colors = rgb.quantize(colors=5).convert("RGB").getcolors(maxcolors=128 * 128) or []
            total = max(1, sum(count for count, _ in colors))
            palette = [
                {"hex": f"#{red:02X}{green:02X}{blue:02X}", "ratio": count / total}
                for count, (red, green, blue) in sorted(colors, reverse=True)[:5]
            ]
            pixels = list(rgb.getdata())
            luminance = sum(0.2126 * red + 0.7152 * green + 0.0722 * blue for red, green, blue in pixels) / max(1, len(pixels))
            saturation = sum(colorsys.rgb_to_hsv(red / 255, green / 255, blue / 255)[1] for red, green, blue in pixels) / max(1, len(pixels))
            return {
                "palette": palette,
                "brightness": "dark" if luminance < 85 else "light" if luminance > 170 else "medium",
                "saturation": "low" if saturation < 0.25 else "high" if saturation > 0.6 else "medium",
            }
    except Exception:
        return None

def _media_type_to_kind(media_type: str) -> str:
    """Map media type to response kind."""
    from sdpm.tools.attachment.limits import IMAGE_MEDIA_TYPES, TEXT_MEDIA_TYPES

    if media_type in IMAGE_MEDIA_TYPES:
        return "image"
    if media_type in TEXT_MEDIA_TYPES:
        return "text"
    if "presentationml" in media_type:
        return "pptx"
    return "document"


def _get_text_projection(
    path: Path,
    media_type: str,
    source_identity: dict[str, Any] | None,
    cache: LocalStageCache | None,
) -> bytes | None:
    """Get or compute the canonical text projection with verified stage caching."""
    if cache is None or source_identity is None:
        return _compute_text_projection(path, media_type)

    source_identity_hash = compute_source_identity_hash(source_identity)
    version = pipeline_version()
    options_hash = compute_hash(json.dumps({"mediaType": media_type}, sort_keys=True))
    pipeline_key = compute_pipeline_key(version, options_hash)
    stage_key = compute_stage_key("extract_text", {"mediaType": media_type})
    record = cache.get_stage(source_identity_hash, pipeline_key, "extract_text", stage_key)
    output_path = cache.stage_dir(source_identity_hash, pipeline_key, "extract_text", stage_key) / "outputs" / "projection.txt"
    if record is not None and output_path.is_file():
        return output_path.read_bytes()

    projection = _compute_text_projection(path, media_type)
    if projection is None:
        return None

    import tempfile
    from datetime import datetime, timezone

    with tempfile.TemporaryDirectory(prefix="sdpm-stage-") as tmp:
        outputs = Path(tmp)
        projection_path = outputs / "projection.txt"
        projection_path.write_bytes(projection)
        record = StageRecord(
            stage="extract_text",
            source_identity_hash=source_identity_hash,
            source_hash=compute_file_sha256(path),
            pipeline_version=version,
            options_hash=options_hash,
            outputs=[{
                "path": "projection.txt",
                "size": len(projection),
                "sha256": compute_hash(projection),
                "contentType": "text/plain",
            }],
            completed_at=datetime.now(timezone.utc).isoformat(),
        )
        cache.publish_stage(
            source_identity_hash, pipeline_key, "extract_text", stage_key, record, outputs,
        )
    return projection


def _compute_text_projection(path: Path, media_type: str) -> bytes | None:
    from sdpm.tools.attachment.limits import TEXT_MEDIA_TYPES

    if media_type in TEXT_MEDIA_TYPES:
        data = path.read_bytes()
        try:
            data.decode("utf-8", errors="strict")
        except UnicodeDecodeError:
            raise SourceValidationError("File is not valid UTF-8")
        return data
    if "presentationml" in media_type:
        return _pptx_projection(path)[0]
    if media_type == "application/pdf":
        return _pdf_text_projection(path)
    if "wordprocessingml" in media_type:
        return _docx_text_projection(path)
    if "spreadsheetml" in media_type:
        return _xlsx_text_projection(path)
    return None

def _pptx_projection(path: Path) -> tuple[bytes, dict[str, Any]]:
    """Create PPTX text summary and edit-workflow header fields."""
    from sdpm.engine.converter import pptx_to_json as _convert
    from sdpm.tools.attachment import PPTX_GUIDE_INSTRUCTION

    result = _convert(path)
    slides = result.get("slides", []) if isinstance(result, dict) else []

    summary_lines = [f"# PPTX Summary: {path.name}", f"Slides: {len(slides)}", ""]

    for i, slide in enumerate(slides, 1):
        title = slide.get("title", f"Slide {i}")
        summary_lines.append(f"## Slide {i}: {title}")
        for elem in slide.get("elements", []):
            text = elem.get("text", "")
            if text:
                summary_lines.append(f"  {text[:200]}")
        summary_lines.append("")

    fonts = result.get("fonts", {}) if isinstance(result, dict) else {}
    header = {
        "guide": "import-pptx",
        "guideInstruction": PPTX_GUIDE_INSTRUCTION,
        "suggestedName": path.stem,
        "slideCount": len(slides),
        "themeHints": {"fonts": fonts, "accentColors": []},
    }
    return "\n".join(summary_lines).encode("utf-8"), header


def _pdf_text_projection(path: Path) -> bytes:
    """Extract text from PDF as Markdown."""
    try:
        import pdfplumber
    except ImportError:
        raise ImportConversionFailed("pdfplumber not available for PDF text extraction")

    lines: list[str] = []
    try:
        with pdfplumber.open(path) as pdf:
            page_count = len(pdf.pages)
            if page_count > MAX_PDF_PAGES:
                lines.append(f"[PDF truncated: showing first {MAX_PDF_PAGES} of {page_count} pages]")
                page_count = MAX_PDF_PAGES

            for i, page in enumerate(pdf.pages[:page_count], 1):
                lines.append(f"--- Page {i} ---")
                text = page.extract_text() or ""
                lines.append(text)

                # Tables
                tables = page.extract_tables()
                for table in tables:
                    if table:
                        lines.append("")
                        for row in table:
                            cells = [str(c or "") for c in row]
                            lines.append("| " + " | ".join(cells) + " |")
                lines.append("")
    except Exception as e:
        raise ImportConversionFailed(f"PDF text extraction failed: {e}")

    return "\n".join(lines).encode("utf-8")


def _docx_text_projection(path: Path) -> bytes:
    """Extract text from DOCX as Markdown."""
    try:
        from docx import Document
    except ImportError:
        raise ImportConversionFailed("python-docx not available for DOCX text extraction")

    lines: list[str] = []
    try:
        doc = Document(str(path))
        for para in doc.paragraphs:
            text = para.text.strip()
            if text:
                # Simple heading detection
                if para.style and para.style.name.startswith("Heading"):
                    level = para.style.name.replace("Heading ", "")
                    try:
                        level_num = int(level)
                    except ValueError:
                        level_num = 1
                    lines.append(f"{'#' * level_num} {text}")
                else:
                    lines.append(text)
            lines.append("")

        # Tables
        for table in doc.tables:
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells]
                lines.append("| " + " | ".join(cells) + " |")
            lines.append("")
    except Exception as e:
        raise ImportConversionFailed(f"DOCX text extraction failed: {e}")

    return "\n".join(lines).encode("utf-8")


def _xlsx_text_projection(path: Path) -> bytes:
    """Extract text from XLSX as Markdown tables."""
    try:
        import openpyxl
    except ImportError:
        raise ImportConversionFailed("openpyxl not available for XLSX text extraction")

    lines: list[str] = []
    try:
        wb = openpyxl.load_workbook(str(path), read_only=True, data_only=True)
        for sheet_name in wb.sheetnames:
            ws = wb[sheet_name]
            lines.append(f"## Sheet: {sheet_name}")
            lines.append("")
            for row in ws.iter_rows(values_only=True):
                cells = [str(c) if c is not None else "" for c in row]
                if any(cells):
                    lines.append("| " + " | ".join(cells) + " |")
            lines.append("")
        wb.close()
    except Exception as e:
        raise ImportConversionFailed(f"XLSX text extraction failed: {e}")

    return "\n".join(lines).encode("utf-8")


def _convert_webp_to_png(src: Path, dest: Path) -> None:
    """Convert WebP to PNG for compatibility."""
    try:
        from PIL import Image
        with Image.open(src) as img:
            img.save(dest, "PNG")
    except Exception as error:
        raise ImportConversionFailed(f"WebP to PNG conversion failed: {error}")


def _import_pptx(
    source_path: Path,
    staging: Path,
    completed_stages: list[str],
    complete_stage: Any,
    check_deadline: Any,
) -> None:
    """Convert a PPTX, resuming from the latest verified cached stage."""
    from sdpm.engine.converter import pptx_to_json as _convert
    from sdpm.engine.converter.template import extract_placeholder_template

    converted = staging / ".converted-pptx"
    deck_dir = staging / "deck"
    if "extract_text" not in completed_stages:
        check_deadline()
        try:
            _convert(source_path, output_dir=converted)
        except Exception as error:
            raise ImportConversionFailed(f"PPTX conversion failed: {error}")
        source_deck_json = converted / "deck.json"
        source_slides = converted / "slides"
        if not source_deck_json.is_file() or not source_slides.is_dir():
            raise ImportConversionFailed("PPTX converter did not produce deck.json and slides/")
        deck_dir.mkdir(exist_ok=True)
        shutil.copy2(source_deck_json, deck_dir / "deck.json")
        shutil.copytree(source_slides, deck_dir / "slides", dirs_exist_ok=True)
        complete_stage("extract_text")

    if "extract_images" not in completed_stages:
        check_deadline()
        converted_images = converted / "images"
        if converted_images.is_dir():
            shutil.copytree(converted_images, staging / "extracted" / "images", dirs_exist_ok=True)
        complete_stage("extract_images")

    if "convert_deck" not in completed_stages:
        check_deadline()
        try:
            deck_dir.mkdir(exist_ok=True)
            extract_placeholder_template(source_path, deck_dir / "template.pptx")
        except Exception as error:
            raise ImportConversionFailed(f"Placeholder template extraction failed: {error}")
        shutil.rmtree(converted, ignore_errors=True)
        complete_stage("convert_deck")

def _import_document(
    source_path: Path,
    media_type: str,
    staging: Path,
    completed_stages: list[str],
    complete_stage: Any,
    check_deadline: Any,
) -> None:
    """Import a document, resuming from verified extracted outputs."""
    from shared.ingest import convert_file

    extracted_dir = staging / "extracted"
    if "extract_text" not in completed_stages:
        check_deadline()
        extracted_dir.mkdir(parents=True, exist_ok=True)
        try:
            result = convert_file(source_path, extracted_dir)
            if result.status == "error":
                raise ImportConversionFailed(result.error or "Document conversion failed")
            if media_type == "application/pdf":
                images_dir = extracted_dir / "images"
                extracted_image_bytes = (
                    sum(file.stat().st_size for file in images_dir.rglob("*") if file.is_file())
                    if images_dir.exists()
                    else 0
                )
                if extracted_image_bytes > MAX_PDF_EXTRACTED_IMAGES_BYTES:
                    raise ImportLimitExceeded("PDF extracted images exceed 200 MiB")
        except (ImportConversionFailed, ImportLimitExceeded):
            raise
        except Exception as error:
            raise ImportConversionFailed(f"Document conversion failed: {error}")
        complete_stage("extract_text")

    if "extract_images" not in completed_stages:
        check_deadline()
        complete_stage("extract_images")
    if "convert_deck" not in completed_stages:
        check_deadline()
        complete_stage("convert_deck")  # N/A for documents


def _build_image_mapping(staging: Path) -> dict[str, str]:
    """Build image mapping from extracted images."""
    mapping: dict[str, str] = {}
    images_dir = staging / "extracted" / "images"
    if images_dir.exists():
        for img in images_dir.iterdir():
            if img.is_file():
                # Map original name → bundle-relative path
                mapping[img.name] = f"extracted/images/{img.name}"
    return mapping


def _find_deck_json(staging: Path) -> str | None:
    """Find deck.json relative path in staging."""
    deck_json = staging / "deck" / "deck.json"
    if deck_json.exists():
        return "deck/deck.json"
    return None


def _find_template(staging: Path) -> str | None:
    """Find template.pptx relative path in staging."""
    template = staging / "deck" / "template.pptx"
    if template.exists():
        return "deck/template.pptx"
    return None
