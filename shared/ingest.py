# Copyright Amazon.com, Inc. or its affiliates. All Rights Reserved.
# SPDX-License-Identifier: MIT-0
"""Upload file conversion pipeline (Cloud/Local shared).

Converts uploaded binary files to agent-readable formats at upload time.
Pure function: takes file path + output dir, writes converted files, returns result.
No I/O dependencies (S3, DynamoDB) — callers handle storage.

Conversion matrix:
    Image/Text  → no conversion (copy as-is by caller)
    PDF         → pdfplumber (text+table+image position) + pypdf (image binary)
    DOCX        → python-docx (text+table interleaved) + zipfile (image binary)
    XLSX        → openpyxl (table structure) + zipfile (image binary)
    PPTX        → pptx_to_json Engine
"""

from __future__ import annotations

import logging
import re
import zipfile
from dataclasses import dataclass, field
from pathlib import Path
from typing import Literal

logger = logging.getLogger(__name__)

# PDF page limit — pages beyond this are skipped with a warning.
_PDF_MAX_PAGES = 100

# File types that need no conversion (the attachment pipeline keeps raw sources).
IMAGE_EXTS = {".png", ".jpg", ".jpeg", ".gif", ".webp", ".svg"}
TEXT_EXTS = {".csv", ".json", ".txt", ".md", ".html"}
PASSTHROUGH_EXTS = IMAGE_EXTS | TEXT_EXTS

# Backward-compat aliases (internal use within this module).
_IMAGE_EXTS = IMAGE_EXTS
_TEXT_EXTS = TEXT_EXTS
_PASSTHROUGH_EXTS = PASSTHROUGH_EXTS


@dataclass
class ConversionResult:
    """Result of a file conversion."""

    status: Literal["success", "partial", "error"]
    markdown: str | None = None
    json_data: str | None = None
    images: list[str] = field(default_factory=list)
    warnings: list[str] = field(default_factory=list)
    error: str | None = None
    # Deck-structure metadata (populated only when output_dir contains
    # deck.json + slides/, i.e. for PPTX after pptx-import-edit).
    deck_structure: bool = False
    slide_count: int = 0
    theme_hints: dict | None = None
    suggested_name: str | None = None
    # Path (relative to output_dir) of the placeholder-only template extracted
    # from the source PPTX. Populated only when deck_structure is True.
    template_path: str | None = None


def convert_file(file_path: Path, output_dir: Path) -> ConversionResult:
    """Convert a file and write results to output_dir.

    - output_dir/{name}.md or output_dir/slides.json
    - output_dir/images/ for extracted images
    - Caller decides what output_dir maps to (S3 prefix, local path, etc.)
    """
    ext = file_path.suffix.lower()

    if ext in _PASSTHROUGH_EXTS:
        return ConversionResult(status="success")

    output_dir.mkdir(parents=True, exist_ok=True)
    images_dir = output_dir / "images"

    if ext == ".pdf":
        return _convert_pdf(file_path, output_dir, images_dir)
    if ext == ".docx":
        return _convert_docx(file_path, output_dir, images_dir)
    if ext == ".xlsx":
        return _convert_xlsx(file_path, output_dir, images_dir)
    if ext == ".pptx":
        return _convert_pptx(file_path, output_dir)

    return ConversionResult(status="error", error=f"Unsupported file type: {ext}")


# ---------------------------------------------------------------------------
# PDF: pdfplumber (text + tables + image positions) + pypdf (image binaries)
# ---------------------------------------------------------------------------


def _convert_pdf(file_path: Path, output_dir: Path, images_dir: Path) -> ConversionResult:
    """Convert PDF with text/table/image interleaving by Y coordinate."""
    try:
        import pdfplumber
        from pypdf import PdfReader
    except ImportError as e:
        return ConversionResult(status="error", error=f"Missing dependency: {e}")

    warnings: list[str] = []
    images: list[str] = []
    md_parts: list[str] = []

    try:
        reader = PdfReader(str(file_path))
    except Exception as e:
        return ConversionResult(status="error", error=f"Cannot open PDF: {e}")

    total_pages = len(reader.pages)
    process_pages = min(total_pages, _PDF_MAX_PAGES)
    if total_pages > _PDF_MAX_PAGES:
        warnings.append(f"PDF has {total_pages} pages; only first {_PDF_MAX_PAGES} processed.")

    try:
        pdf = pdfplumber.open(str(file_path))
    except Exception as e:
        return ConversionResult(status="error", error=f"pdfplumber cannot open PDF: {e}")

    # Document-level decode cache: XObject indirect-object id → written filename.
    # Some producers (e.g. LibreOffice PDF export) put every image of the
    # document into a resource dict shared by all pages; without this cache
    # each page would re-decode every image (pages × images decodes).
    image_cache: dict[object, str] = {}

    for page_idx in range(process_pages):
        page_num = page_idx + 1
        pb_page = pdf.pages[page_idx]
        pypdf_page = reader.pages[page_idx]

        # Collect content fragments with Y positions for interleaving
        fragments: list[tuple[float, str]] = []

        # --- Text lines ---
        try:
            lines = pb_page.extract_text_lines() or []
            for line in lines:
                fragments.append((line["top"], line["text"]))
        except Exception:
            pass

        # --- Tables ---
        try:
            tables = pb_page.find_tables() or []
            # Track table bounding boxes to exclude overlapping text
            table_bboxes = []
            for tbl in tables:
                bbox = tbl.bbox  # (x0, top, x1, bottom)
                table_bboxes.append(bbox)
                rows = tbl.extract() or []
                if rows:
                    md_table = _rows_to_md_table(rows)
                    fragments.append((bbox[1], md_table))

            # Remove text lines that fall inside table bounding boxes
            if table_bboxes:
                fragments = [
                    (y, text) for y, text in fragments
                    if not any(
                        bbox[1] <= y <= bbox[3] and not text.startswith("|")
                        for bbox in table_bboxes
                    )
                ]
        except Exception:
            pass

        # --- Images: positions from pdfplumber (actually drawn on the page),
        # binaries from pypdf, matched by resource name. Decode results are
        # cached document-wide by indirect-object id, so images reused across
        # pages (shared resource dicts, repeated logos) are decoded once. ---
        try:
            pb_images = pb_page.images or []
            if pb_images:
                images_dir.mkdir(parents=True, exist_ok=True)
            for pb_img in pb_images:
                res_name = pb_img.get("name")  # e.g. "Im117"
                y_pos = pb_img.get("top", float("inf"))
                if not res_name:
                    warnings.append(
                        f"Page {page_num}: unnamed (inline) image skipped."
                    )
                    continue
                key = f"/{res_name}"
                try:
                    xobjects = pypdf_page["/Resources"]["/XObject"].get_object()
                    ref = xobjects.raw_get(key)
                    cache_key = getattr(ref, "idnum", None) or (page_num, key)
                except Exception:
                    cache_key = (page_num, key)

                img_name = image_cache.get(cache_key)
                if img_name is None:
                    try:
                        pypdf_img = pypdf_page.images[key]
                    except Exception as e:
                        warnings.append(
                            f"Page {page_num}: image {res_name} extraction failed: {e}"
                        )
                        continue
                    ext = pypdf_img.name.rsplit(".", 1)[-1] if "." in pypdf_img.name else "png"
                    img_name = f"pdf_p{page_num}_{res_name}.{ext}"
                    (images_dir / img_name).write_bytes(pypdf_img.data)
                    images.append(img_name)
                    image_cache[cache_key] = img_name
                fragments.append((y_pos, f"![{img_name}]({img_name})"))
        except Exception as e:
            warnings.append(f"Page {page_num}: image extraction failed: {e}")

        # Sort by Y position and build page markdown
        fragments.sort(key=lambda f: f[0])
        page_text = "\n\n".join(text for _, text in fragments)
        if page_text.strip():
            md_parts.append(f"### Page {page_num}\n\n{page_text}")

    pdf.close()

    markdown = "\n\n".join(md_parts)
    if not markdown.strip() and not images:
        return ConversionResult(
            status="error", error="No extractable content in PDF."
        )

    name = file_path.stem
    (output_dir / f"{name}.md").write_text(markdown, encoding="utf-8")

    status: Literal["success", "partial"] = "success"
    if not markdown.strip() and images:
        warnings.append("No text extracted (possibly a scanned PDF). Images were extracted.")
        status = "partial"
    elif warnings:
        status = "partial"

    return ConversionResult(
        status=status, markdown=markdown, images=images, warnings=warnings,
    )


def _rows_to_md_table(rows: list[list]) -> str:
    """Convert table rows to Markdown table string."""
    if not rows:
        return ""
    # Sanitize cells
    clean = []
    for row in rows:
        clean.append([(cell or "").replace("|", "\\|").replace("\n", " ") for cell in row])

    header = "| " + " | ".join(clean[0]) + " |"
    sep = "| " + " | ".join("---" for _ in clean[0]) + " |"
    body = "\n".join("| " + " | ".join(r) + " |" for r in clean[1:])
    return f"{header}\n{sep}\n{body}" if body else f"{header}\n{sep}"


# ---------------------------------------------------------------------------
# DOCX: python-docx (text + tables interleaved) + zipfile (image binaries)
# ---------------------------------------------------------------------------


def _convert_docx(file_path: Path, output_dir: Path, images_dir: Path) -> ConversionResult:
    """Convert DOCX to Markdown with image extraction.

    Iterates body elements in document order so paragraphs and tables are
    interleaved correctly (unlike doc.paragraphs + doc.tables which are separate).
    """
    try:
        from docx import Document
        from docx.table import Table
        from docx.text.paragraph import Paragraph
    except ImportError as e:
        return ConversionResult(status="error", error=f"Missing dependency: {e}")

    warnings: list[str] = []
    images: list[str] = []

    try:
        doc = Document(str(file_path))
    except Exception as e:
        return ConversionResult(status="error", error=f"Cannot open DOCX: {e}")

    lines: list[str] = []
    num_counter = 0

    for element in doc.element.body:
        tag = element.tag.split("}")[-1]

        if tag == "p":
            para = Paragraph(element, doc)
            style = para.style.name
            text = para.text

            if style == "Title":
                num_counter = 0
                lines.append(f"# {text}")
                lines.append("")
            elif style.startswith("Heading"):
                num_counter = 0
                level = int(style.split()[-1]) if style[-1].isdigit() else 1
                lines.append(f"{'#' * level} {text}")
                lines.append("")
            elif style == "List Bullet":
                lines.append(f"- {text}")
            elif style == "List Number":
                num_counter += 1
                lines.append(f"{num_counter}. {text}")
            elif text.strip():
                num_counter = 0
                lines.append(text)
                lines.append("")
            else:
                num_counter = 0

        elif tag == "tbl":
            table = Table(element, doc)
            rows: list[list[str]] = []
            for row in table.rows:
                cells = [cell.text.strip().replace("|", "\\|").replace("\n", " ") for cell in row.cells]
                rows.append(cells)
            if rows:
                lines.append(_rows_to_md_table(rows))
                lines.append("")

    # Image extraction via zipfile
    try:
        with zipfile.ZipFile(str(file_path), "r") as zf:
            media_files = [n for n in zf.namelist() if n.startswith("word/media/")]
            if media_files:
                images_dir.mkdir(parents=True, exist_ok=True)
                for media_path in media_files:
                    img_name = Path(media_path).name
                    data = zf.read(media_path)
                    (images_dir / img_name).write_bytes(data)
                    images.append(img_name)
    except Exception as e:
        warnings.append(f"Image extraction failed: {e}")

    markdown = "\n".join(lines)

    if not markdown.strip() and not images:
        return ConversionResult(status="error", error="No extractable content in DOCX.")

    name = file_path.stem
    (output_dir / f"{name}.md").write_text(markdown, encoding="utf-8")

    status: Literal["success", "partial"] = "partial" if warnings else "success"
    return ConversionResult(
        status=status, markdown=markdown, images=images, warnings=warnings,
    )


# ---------------------------------------------------------------------------
# XLSX: openpyxl (table structure) + zipfile (image binaries)
# ---------------------------------------------------------------------------


def _convert_xlsx(file_path: Path, output_dir: Path, images_dir: Path) -> ConversionResult:
    """Convert XLSX to Markdown tables with image extraction."""
    try:
        from openpyxl import load_workbook
    except ImportError as e:
        return ConversionResult(status="error", error=f"Missing dependency: {e}")

    warnings: list[str] = []
    images: list[str] = []

    try:
        wb = load_workbook(str(file_path), data_only=True, read_only=True)
    except Exception as e:
        return ConversionResult(status="error", error=f"Cannot open XLSX: {e}")

    md_parts: list[str] = []
    try:
        for ws in wb.worksheets:
            rows = list(ws.iter_rows(values_only=True))
            if not rows:
                continue
            lines: list[str] = [f"## {ws.title}"]
            # Header
            header = [str(c) if c is not None else "" for c in rows[0]]
            lines.append("| " + " | ".join(header) + " |")
            lines.append("| " + " | ".join("---" for _ in header) + " |")
            # Data rows
            for row in rows[1:]:
                cells = [str(c) if c is not None else "" for c in row]
                lines.append("| " + " | ".join(cells) + " |")
            md_parts.append("\n".join(lines))
    finally:
        wb.close()

    # Image extraction via zipfile
    try:
        with zipfile.ZipFile(str(file_path), "r") as zf:
            media_files = [n for n in zf.namelist() if n.startswith("xl/media/")]
            if media_files:
                images_dir.mkdir(parents=True, exist_ok=True)
                for media_path in media_files:
                    img_name = Path(media_path).name
                    data = zf.read(media_path)
                    (images_dir / img_name).write_bytes(data)
                    images.append(img_name)
    except Exception as e:
        warnings.append(f"Image extraction failed: {e}")

    markdown = "\n\n".join(md_parts)

    if not markdown.strip() and not images:
        return ConversionResult(status="error", error="No extractable content in XLSX.")

    name = file_path.stem
    (output_dir / f"{name}.md").write_text(markdown, encoding="utf-8")

    status: Literal["success", "partial"] = "partial" if warnings else "success"
    return ConversionResult(
        status=status, markdown=markdown, images=images, warnings=warnings,
    )


# ---------------------------------------------------------------------------
# PPTX: pptx_to_json Engine + deck-structure metadata extraction
# ---------------------------------------------------------------------------


def _extract_theme_hints(pptx_path: Path, deck_json_path: Path, slides_dir: Path) -> dict:
    """Extract theme hints (backgroundLuminance, accentColors, fonts) from a converted deck.

    Keys are stable across Local / Cloud / DynamoDB: ``backgroundLuminance``,
    ``accentColors``, ``fonts``.

    Background luminance strategy (prioritised):
      1. If a slide declares an explicit ``background`` (hex), use that slide's color.
      2. Else, fall back to the template's light theme color (lt1 if available,
         otherwise dk1).
      3. Compute per-slide luminance (0.299 R + 0.587 G + 0.114 B, 0-1) and
         return the median across all slides.
    """
    import json as _json
    import statistics

    from sdpm.engine.converter.color import extract_theme_colors_and_mapping

    # Theme colors. PowerPoint can ship multiple slideMasters with
    # different clrMaps; corporate decks frequently flip bg1=dk1 (dark
    # deck) on a secondary master while keeping master 0 light. Walk all
    # masters and pick the one whose clrMap most slides actually use, so
    # backgroundLuminance reflects the rendered deck rather than master 0.
    theme_colors: dict = {}
    color_mapping: dict = {}
    try:
        # Default to master 0.
        theme_colors, color_mapping, _ = extract_theme_colors_and_mapping(pptx_path, 0)
    except Exception:
        pass

    # Discover all slideMasters and pick the one used by the most slides.
    try:
        with zipfile.ZipFile(pptx_path) as zf:
            names = zf.namelist()
            masters = sorted(
                n for n in names
                if n.startswith("ppt/slideMasters/slideMaster") and n.endswith(".xml")
            )
            # layout → master index (0-based per master ordering above)
            layout_to_master: dict[str, int] = {}
            for idx, master in enumerate(masters):
                rels_name = master.replace("slideMasters/", "slideMasters/_rels/") + ".rels"
                if rels_name not in names:
                    continue
                rels_xml = zf.read(rels_name).decode("utf-8", errors="ignore")
                for layout_match in re.finditer(r'Target="\.\./slideLayouts/(slideLayout\d+\.xml)"', rels_xml):
                    layout_to_master[layout_match.group(1)] = idx

            # slide → layout → master
            master_usage: dict[int, int] = {}
            slide_files = sorted(
                n for n in names
                if re.match(r"ppt/slides/slide\d+\.xml$", n)
            )
            for slide in slide_files:
                rels_name = slide.replace("slides/", "slides/_rels/") + ".rels"
                if rels_name not in names:
                    continue
                rels_xml = zf.read(rels_name).decode("utf-8", errors="ignore")
                m = re.search(r'Target="\.\./slideLayouts/(slideLayout\d+\.xml)"', rels_xml)
                if m:
                    layout = m.group(1)
                    master_idx = layout_to_master.get(layout)
                    if master_idx is not None:
                        master_usage[master_idx] = master_usage.get(master_idx, 0) + 1

            if master_usage:
                dominant_master = max(master_usage.items(), key=lambda kv: kv[1])[0]
                if dominant_master != 0:
                    try:
                        theme_colors, color_mapping, _ = extract_theme_colors_and_mapping(
                            pptx_path, dominant_master
                        )
                    except Exception:
                        pass
    except Exception:
        # Multi-master discovery is best-effort; fall through to master 0
        # values that were already loaded above.
        pass

    # Resolve bg1 through clrMap. PowerPoint's default mapping is
    # bg1=lt1 / tx1=dk1, but corporate decks frequently flip it
    # (bg1=dk1 / tx1=lt1) to declare a dark theme without changing the
    # theme XML. The dominant slideMaster's clrMap is the authoritative
    # answer.
    bg1_target = color_mapping.get("bg1", "lt1") if color_mapping else "lt1"
    resolved_bg = theme_colors.get(bg1_target)
    lt1 = theme_colors.get("lt1", "#FFFFFF")
    dk1 = theme_colors.get("dk1", "#000000")
    # Fallback chain: slide bg > clrMap-resolved bg1 > lt1 > dk1 inverted
    default_bg = resolved_bg if resolved_bg else (lt1 if lt1 else _invert_hex(dk1))

    def _hex_to_luminance(hex_color: str) -> float:
        h = hex_color.lstrip("#")
        try:
            r = int(h[0:2], 16)
            g = int(h[2:4], 16)
            b = int(h[4:6], 16)
        except (ValueError, IndexError):
            return 0.5
        return (0.299 * r + 0.587 * g + 0.114 * b) / 255.0

    luminances: list[float] = []
    if slides_dir.is_dir():
        for slide_file in sorted(slides_dir.glob("slide-*.json")):
            try:
                data = _json.loads(slide_file.read_text(encoding="utf-8"))
            except Exception:
                continue
            bg = data.get("background") or default_bg
            luminances.append(_hex_to_luminance(bg))
    if not luminances:
        luminances = [_hex_to_luminance(default_bg)]

    background_luminance = statistics.median(luminances)

    accent_colors: list[str] = []
    for name in ("accent1", "accent2", "accent3"):
        c = theme_colors.get(name)
        if isinstance(c, str) and c.startswith("#"):
            accent_colors.append(c.upper())
    accent_colors = accent_colors[:3]

    fonts: dict = {}
    if deck_json_path.exists():
        try:
            deck = _json.loads(deck_json_path.read_text(encoding="utf-8"))
            f = deck.get("fonts")
            if isinstance(f, dict):
                fonts = f
        except Exception:
            pass

    return {
        "backgroundLuminance": round(background_luminance, 4),
        "accentColors": accent_colors,
        "fonts": fonts,
    }


def _invert_hex(hex_color: str) -> str:
    """Invert a #RRGGBB color (for the dk1-inverted fallback)."""
    h = hex_color.lstrip("#")
    try:
        r = 255 - int(h[0:2], 16)
        g = 255 - int(h[2:4], 16)
        b = 255 - int(h[4:6], 16)
    except (ValueError, IndexError):
        return "#FFFFFF"
    return f"#{r:02X}{g:02X}{b:02X}"


def _convert_pptx(file_path: Path, output_dir: Path) -> ConversionResult:
    """Convert PPTX to deck structure via Engine's pptx_to_json.

    Populates deck_structure, slide_count, theme_hints, suggested_name when
    the output directory contains the new deck layout (deck.json + slides/).
    """
    import json

    try:
        from sdpm.engine.converter import pptx_to_json
    except ImportError as e:
        return ConversionResult(status="error", error=f"Missing dependency: {e}")

    try:
        result = pptx_to_json(file_path, output_dir=output_dir)
        json_str = json.dumps(result, ensure_ascii=False)

        # Collect extracted images
        images: list[str] = []
        img_dir = output_dir / "images"
        if img_dir.exists():
            images = [f.name for f in img_dir.iterdir() if f.is_file()]

        deck_json_path = output_dir / "deck.json"
        slides_dir = output_dir / "slides"
        deck_structure = deck_json_path.exists() and slides_dir.is_dir()

        slide_count = 0
        theme_hints: dict | None = None
        suggested_name: str | None = None
        template_path: str | None = None
        warnings: list[str] = []
        if deck_structure:
            slide_count = len(list(slides_dir.glob("slide-*.json")))
            try:
                theme_hints = _extract_theme_hints(file_path, deck_json_path, slides_dir)
            except Exception as e:
                logger.warning("theme_hints extraction failed: %s", e)
            suggested_name = file_path.stem
            try:
                from sdpm.engine.converter.template import extract_placeholder_template

                template_out = output_dir / "template.pptx"
                extract_placeholder_template(file_path, template_out)
                template_path = "template.pptx"
            except Exception as e:
                logger.warning("placeholder template extraction failed: %s", e)
                warnings.append(f"placeholder template extraction failed: {e}")

        return ConversionResult(
            status="success",
            json_data=json_str,
            images=images,
            warnings=warnings,
            deck_structure=deck_structure,
            slide_count=slide_count,
            theme_hints=theme_hints,
            suggested_name=suggested_name,
            template_path=template_path,
        )
    except Exception as e:
        return ConversionResult(status="error", error=f"PPTX conversion failed: {e}")
